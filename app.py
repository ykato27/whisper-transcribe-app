"""
Whisper文字起こし + Gemini議事録生成アプリ（完全版 + Word/VTT対応）
"""

import os
import tempfile
import whisper
import torch
import streamlit as st
import subprocess
import time
from pydub import AudioSegment
import math
import google.generativeai as genai
from datetime import datetime
import json
import re

# Word/VTT用ライブラリ
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import webvtt
    WEBVTT_AVAILABLE = True
except ImportError:
    WEBVTT_AVAILABLE = False

# ページ設定
st.set_page_config(page_title="AI議事録作成ツール", page_icon="📝", layout="wide")

# セッション状態の初期化
if "transcribed_text" not in st.session_state:
    st.session_state.transcribed_text = ""
if "minutes" not in st.session_state:
    st.session_state.minutes = ""
if "file_type" not in st.session_state:
    st.session_state.file_type = None
if "api_key" not in st.session_state:
    st.session_state.api_key = ""
if "gemini_model" not in st.session_state:
    st.session_state.gemini_model = "gemini-2.5-flash"
if "custom_prompts" not in st.session_state:
    st.session_state.custom_prompts = {}


# キャッシュ設定
@st.cache_resource
def load_whisper_model(model_name):
    """Whisperモデルをロード（キャッシュ）"""
    device = "cuda" if torch.cuda.is_available() else "cpu"
    return whisper.load_model(model_name, device=device)


def check_ffmpeg():
    """FFmpegの存在確認"""
    try:
        result = subprocess.run(
            ["ffmpeg", "-version"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=5,
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def get_available_models():
    """利用可能なモデルリスト"""
    return ["tiny", "base", "small", "medium"]


def process_audio_chunk(model, audio_segment, language=None):
    """音声チャンクを処理"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as chunk_file:
        try:
            audio_segment.export(chunk_file.name, format="wav")
            options = {"language": language} if language else {}
            result = model.transcribe(chunk_file.name, **options)
            return result["text"]
        finally:
            try:
                os.unlink(chunk_file.name)
            except:
                pass


def get_default_prompt_templates():
    """デフォルトの議事録生成用プロンプトテンプレート"""
    return {
        "標準ビジネス議事録": """# 役割
あなたは日本企業で働く経験豊富なビジネスコンサルタントです。会議の文字起こしテキストから日本語でプロフェッショナルな議事録を作成します。

# 入力情報
- 日時: {date}
- 文字起こしテキスト:
{transcript}

# 出力形式
以下の形式で日本語の議事録を作成してください:

## 議事録

### 会議情報
* 日時: {date}
* 参加者
   * 先方: [文字起こしから判断して記載、不明な場合は「[記載なし]」]
   * 自社: [文字起こしから判断して記載、不明な場合は「[記載なし]」]

### AR（Action Required/宿題事項）
* 先方:
   * [具体的なアクション項目] - 担当: [名前] - 期限: [日付]
   * [決定していない場合は「[宿題事項なし]」]
* 自社:
   * [具体的なアクション項目] - 担当: [名前] - 期限: [日付]
   * [決定していない場合は「[宿題事項なし]」]

### 決定事項
* [明確に決定した事項のみを箇条書き]
* [決定事項がない場合は「[決定事項なし]」]

### 議事メモ
* アジェンダ①: [トピック名]
   * [詳細な議論内容を記載]
   * [発言の背景・理由・懸念事項も含める]
   * [数値やデータは正確に転記]
* アジェンダ②: [トピック名]
   * [詳細な議論内容を記載]
* アジェンダ③: [トピック名]
   * [詳細な議論内容を記載]

# 作成時の厳守事項

## 1. 言語
- **必ず日本語で作成**: すべての内容を日本語で記載
- 英語の文字起こしの場合も、日本語に翻訳して議事録を作成

## 2. 憶測・推測の完全禁止
- 文字起こしに明記されていない情報は一切追加しない
- 不明確な点は「[記載なし]」または「[要確認]」と明記
- 文脈から推測できそうな内容でも、明示されていなければ記載しない

## 3. 参加者
- 文字起こしに記載されている通りに正確に転記
- 部署名・役職が不明な場合は名前のみ記載
- 先方と自社を明確に区別
- 判断できない場合は「[記載なし]」

## 4. AR（宿題事項）
- 文字起こしに明示されているアクション項目のみを記載
- 担当者名が明記されている場合のみ記載（不明な場合は「[担当者未定]」）
- 期限が示されている場合は必ず記載（不明な場合は「[期限未定]」）
- 曖昧な発言は宿題事項に含めず、議事メモに記載
- 先方と自社のARを明確に分けて記載

## 5. 決定事項
- 明確に「決定した」「合意した」と分かる内容のみ記載
- 「検討する」「相談する」等の未確定事項は決定事項に含めない
- 決定事項が文字起こしから読み取れない場合は「[決定事項なし]」と記載

## 6. 議事メモ
- **網羅性を最優先**: 文字起こしの内容を可能な限り詳細に記載
- 議事メモは多少長くなっても構わないので、重要な情報を漏らさない
- 発言の背景・理由・懸念事項なども含めて記載
- 数値・データ・固有名詞は一字一句正確に転記
- 議論の流れや文脈が理解できるよう丁寧に記述
- 些細に見える発言でも、後で重要になる可能性があるため記載
- 未解決の論点や持ち越し事項も必ず記載
- アジェンダごとに整理して記載

# 記載の優先順位
1. 第1優先: 正確性（憶測ゼロ、文字起こしに書かれていることのみ）
2. 第2優先: 網羅性（重要な情報の抜け漏れ防止）
3. 第3優先: 簡潔性（ただし網羅性を犠牲にしない）

# 禁止事項
- ❌ 文字起こしにない情報の補完
- ❌ 「おそらく」「と思われる」等の推測表現
- ❌ 一般的な知識や常識に基づく補足
- ❌ 重要そうな情報の省略や要約のしすぎ
- ❌ 英語での出力（必ず日本語で作成）

上記の形式で議事録を作成してください。
"""
    }


def get_all_prompt_templates():
    """デフォルト + カスタムプロンプトを統合"""
    templates = get_default_prompt_templates()
    templates.update(st.session_state.custom_prompts)
    return templates


def generate_minutes_with_gemini(transcript, prompt_template, api_key, model_name=None):
    """Gemini APIで議事録を生成"""
    try:
        genai.configure(api_key=api_key)

        # モデル名を取得
        if model_name is None:
            model_name = st.session_state.get("gemini_model", "gemini-pro")

        model = genai.GenerativeModel(model_name)

        current_date = datetime.now().strftime("%Y年%m月%d日")
        prompt = prompt_template.format(transcript=transcript, date=current_date)

        response = model.generate_content(prompt)
        return response.text

    except Exception as e:
        raise Exception(f"Gemini API エラー: {str(e)}")


def is_audio_or_video_file(filename):
    """音声・動画ファイルかどうかを判定"""
    audio_video_extensions = [
        ".mp3",
        ".wav",
        ".m4a",
        ".ogg",
        ".flac",
        ".mp4",
        ".avi",
        ".mov",
        ".mkv",
    ]
    ext = os.path.splitext(filename)[1].lower()
    return ext in audio_video_extensions


def is_text_file(filename):
    """テキストファイルかどうかを判定（VTTを含む）"""
    text_extensions = [".txt", ".md", ".text", ".vtt"]
    ext = os.path.splitext(filename)[1].lower()
    return ext in text_extensions


def is_word_file(filename):
    """Wordファイルかどうかを判定"""
    word_extensions = [".docx", ".doc"]
    ext = os.path.splitext(filename)[1].lower()
    return ext in word_extensions


def is_vtt_file(filename):
    """VTTファイルかどうかを判定"""
    ext = os.path.splitext(filename)[1].lower()
    return ext == ".vtt"


def read_word_file(uploaded_file):
    """Wordファイルからテキストを抽出（uploaded_fileを直接使用）"""
    if not DOCX_AVAILABLE:
        raise Exception("python-docxライブラリがインストールされていません。'pip install python-docx'を実行してください。")
    
    try:
        # uploaded_fileを直接Documentに渡す
        doc = Document(uploaded_file)
        
        # 全段落のテキストを結合
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 空行をスキップ
                text_content.append(paragraph.text)
        
        # テーブル内のテキストも抽出
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        return "\n".join(text_content)
    
    except Exception as e:
        raise Exception(f"Wordファイルの読み込みエラー: {str(e)}")



def read_vtt_file(uploaded_file):
    """VTTファイルからテキストを抽出（webvtt-py使用）"""
    try:
        # VTTファイルをデコード
        content = uploaded_file.read().decode("utf-8")
        
        if WEBVTT_AVAILABLE:
            # webvtt-pyを使用して解析
            try:
                vtt_data = webvtt.read_buffer(content.splitlines())
                text_lines = []
                
                for caption in vtt_data:
                    # タグを削除してテキストのみ抽出
                    clean_text = re.sub(r'<[^>]+>', '', caption.text)
                    if clean_text.strip():
                        text_lines.append(clean_text.strip())
                
                return " ".join(text_lines)
            except Exception as e:
                # webvtt-pyで失敗した場合は従来の方法にフォールバック
                st.warning(f"webvtt-pyでの解析に失敗したため、通常の方法で処理します: {str(e)}")
        
        # 従来の方法（webvtt-py未使用またはエラー時）
        if not content.startswith("WEBVTT"):
            raise Exception("有効なWebVTTファイルではありません")
        
        text_lines = []
        lines = content.split("\n")
        
        in_cue = False
        for line in lines:
            line = line.strip()
            
            if "-->" in line:
                in_cue = True
                continue
            
            if not line:
                in_cue = False
                continue
            
            if line.isdigit():
                continue
            
            if line.startswith(("WEBVTT", "NOTE", "STYLE", "REGION")):
                continue
            
            if in_cue:
                clean_line = re.sub(r'<[^>]+>', '', line)
                if clean_line:
                    text_lines.append(clean_line)
        
        return " ".join(text_lines)
    
    except Exception as e:
        raise Exception(f"VTTファイルの読み込みエラー: {str(e)}")



def transcribe_audio(uploaded_file, model_option, language_option):
    """音声ファイルの文字起こし処理"""

    file_size_mb = uploaded_file.size / (1024 * 1024)

    st.info(f"📄 **{uploaded_file.name}** ({file_size_mb:.2f} MB)")

    if file_size_mb > 200:
        st.error("ファイルサイズが200MBを超えています")
        return None

    # 音声プレビュー
    try:
        file_ext = uploaded_file.name.split(".")[-1].lower()
        if file_ext in ["mp3", "wav", "m4a", "ogg"]:
            st.audio(uploaded_file, format=f"audio/{file_ext}")
    except:
        st.warning("プレビューを表示できません")

    # 一時ファイルとして保存
    with tempfile.NamedTemporaryFile(
        delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}"
    ) as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        temp_filename = tmp_file.name

    try:
        # モデルロード
        with st.spinner("🔄 モデルをロード中..."):
            load_start = time.time()
            model = load_whisper_model(model_option)
            load_time = time.time() - load_start
            st.success(f"✅ モデルロード完了（{load_time:.2f}秒）")

        # 音声読み込み
        with st.spinner("📊 音声ファイルを解析中..."):
            audio = AudioSegment.from_file(temp_filename)
            duration_sec = audio.duration_seconds
            st.info(f"⏱️ 音声長: {duration_sec:.1f}秒")

        # チャンク分割設定（10秒単位）
        chunk_size_sec = 10
        num_chunks = math.ceil(duration_sec / chunk_size_sec)

        # 進捗バー
        progress_bar = st.progress(0)
        status_text = st.empty()
        transcribed_text = ""
        start_time = time.time()

        # チャンクごとに処理
        for i in range(num_chunks):
            start_ms = i * chunk_size_sec * 1000
            end_ms = min((i + 1) * chunk_size_sec * 1000, len(audio))
            chunk_audio = audio[start_ms:end_ms]

            # 文字起こし実行
            chunk_text = process_audio_chunk(
                model,
                chunk_audio,
                language=language_option if language_option else None,
            )
            transcribed_text += chunk_text + " "

            # 進捗更新
            progress = (i + 1) / num_chunks
            elapsed = time.time() - start_time
            eta = (elapsed / progress - elapsed) if progress > 0 else 0

            progress_bar.progress(progress)
            status_text.text(
                f"📝 処理中: {progress*100:.1f}% | " f"残り予想時間: {eta:.1f}秒"
            )

        # 完了
        total_time = time.time() - start_time
        progress_bar.progress(1.0)
        status_text.empty()

        st.success(f"🎉 文字起こし完了！（処理時間: {total_time:.2f}秒）")

        return transcribed_text.strip()

    except Exception as e:
        st.error(f"❌ エラーが発生しました: {str(e)}")
        st.exception(e)
        return None

    finally:
        try:
            os.unlink(temp_filename)
        except:
            pass


def show_settings_page():
    """設定ページを表示"""
    st.title("⚙️ 設定")

    # Gemini API設定
    st.markdown("## 🤖 Gemini API設定")

    # モデル選択
    st.markdown("### モデル選択")
    gemini_models = [
        "gemini-2.5-flash",  # 最新の高速・多機能モデル (推奨)
        "gemini-2.5-pro",  # 最新の最高性能モデル
        "gemini-1.5-flash",  # 前世代の高速モデル
        "gemini-1.5-pro",  # 前世代の最高性能モデル
        "gemini-pro",  # 旧世代の安定版
    ]

    model_descriptions = {
        "gemini-2.5-flash": "最新・最速のフラッシュモデル（高性能＆低コスト）",
        "gemini-2.5-pro": "最新・最高性能のプロモデル（複雑なタスク向け）",
        "gemini-1.5-flash": "前世代の高速モデル",
        "gemini-1.5-pro": "前世代の最高性能モデル",
        "gemini-pro": "旧世代の安定版（レガシー）",
    }

    # 現在のモデルがリストにない場合のフォールバックインデックスを設定
    default_index = 0
    try:
        default_index = gemini_models.index(st.session_state.gemini_model)
    except (ValueError, AttributeError):
        default_index = 0

    selected_model = st.selectbox(
        "使用するGeminiモデル",
        gemini_models,
        index=default_index,
        format_func=lambda x: f"{x} - {model_descriptions.get(x, '')}",
        help="gemini-2.5-flash が現在推奨される高速モデルです",
    )

    if st.button("モデルを保存", type="secondary"):
        st.session_state.gemini_model = selected_model
        st.success(f"✅ モデルを「{selected_model}」に設定しました")

    # API Key設定
    st.markdown("---")
    st.markdown("### 🔑 API Key")
    st.markdown(
        """
    Gemini APIを使用するには、API Keyが必要です。  
    [Google AI Studio](https://aistudio.google.com/app/apikey) から無料で取得できます（無料枠あり）。
    """
    )

    api_key_input = st.text_input(
        "API Keyを入力",
        value=st.session_state.api_key,
        type="password",
        help="入力したAPI Keyはセッション中のみ保持されます",
    )

    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("💾 保存", type="primary"):
            st.session_state.api_key = api_key_input
            st.success("✅ API Keyを保存しました")
    with col2:
        if st.button("🗑️ クリア"):
            st.session_state.api_key = ""
            st.info("API Keyをクリアしました")
            st.rerun()

    if st.session_state.api_key:
        st.success("✅ API Key設定済み")
    else:
        st.warning("⚠️ API Keyが未設定です")

    # カスタムプロンプト設定
    st.markdown("---")
    st.markdown("## 📝 カスタムプロンプト管理")

    st.markdown(
        """
    独自の議事録フォーマットを作成できます。  
    `{transcript}` と `{date}` を使用できます。
    """
    )

    # 新規プロンプト追加
    with st.expander("➕ 新しいプロンプトを追加"):
        new_prompt_name = st.text_input("プロンプト名", key="new_prompt_name")
        new_prompt_content = st.text_area(
            "プロンプト内容",
            height=300,
            placeholder="""例:
以下の文字起こしから、技術ミーティングの議事録を作成してください。

# 文字起こし
{transcript}

# 要件
- 技術的な決定事項を明確に
- アクションアイテムを整理
- 次回の議題を抽出

日時: {date}
""",
            key="new_prompt_content",
        )

        if st.button("追加", type="primary"):
            if new_prompt_name and new_prompt_content:
                if new_prompt_name in get_default_prompt_templates():
                    st.error("❌ デフォルトのプロンプト名は使用できません")
                elif new_prompt_name in st.session_state.custom_prompts:
                    st.error("❌ 同じ名前のプロンプトが既に存在します")
                else:
                    st.session_state.custom_prompts[new_prompt_name] = (
                        new_prompt_content
                    )
                    st.success(f"✅ プロンプト「{new_prompt_name}」を追加しました")
                    st.rerun()
            else:
                st.error("プロンプト名と内容を入力してください")

    # 既存のカスタムプロンプト表示
    if st.session_state.custom_prompts:
        st.markdown("### 📋 保存済みカスタムプロンプト")

        for name, content in st.session_state.custom_prompts.items():
            with st.expander(f"📄 {name}"):
                st.code(content, language="markdown")

                col1, col2 = st.columns([1, 1])
                with col1:
                    # 編集機能
                    edit_key = f"edit_{name}"
                    if st.button(f"✏️ 編集", key=f"btn_edit_{name}"):
                        st.session_state[f"editing_{name}"] = True

                with col2:
                    # 削除機能
                    if st.button(f"🗑️ 削除", key=f"btn_delete_{name}"):
                        del st.session_state.custom_prompts[name]
                        st.success(f"削除しました: {name}")
                        st.rerun()

                # 編集モード
                if st.session_state.get(f"editing_{name}", False):
                    edited_content = st.text_area(
                        "編集", value=content, height=300, key=f"edit_area_{name}"
                    )

                    col1, col2 = st.columns([1, 1])
                    with col1:
                        if st.button("💾 保存", key=f"save_{name}"):
                            st.session_state.custom_prompts[name] = edited_content
                            st.session_state[f"editing_{name}"] = False
                            st.success("保存しました")
                            st.rerun()
                    with col2:
                        if st.button("❌ キャンセル", key=f"cancel_{name}"):
                            st.session_state[f"editing_{name}"] = False
                            st.rerun()
    else:
        st.info("カスタムプロンプトはまだ追加されていません")

    # プロンプトのインポート/エクスポート
    st.markdown("---")
    st.markdown("## 📦 インポート/エクスポート")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### 📥 エクスポート")
        if st.session_state.custom_prompts:
            export_data = json.dumps(
                st.session_state.custom_prompts, ensure_ascii=False, indent=2
            )
            st.download_button(
                label="カスタムプロンプトをダウンロード",
                data=export_data,
                file_name="custom_prompts.json",
                mime="application/json",
                use_container_width=True,
            )
        else:
            st.info("エクスポートするプロンプトがありません")

    with col2:
        st.markdown("### 📤 インポート")
        uploaded_json = st.file_uploader(
            "JSONファイルをアップロード",
            type=["json"],
            help="以前エクスポートしたプロンプトをインポート",
        )

        if uploaded_json:
            try:
                imported_data = json.loads(uploaded_json.read())
                if st.button("インポート実行", type="primary"):
                    st.session_state.custom_prompts.update(imported_data)
                    st.success(
                        f"✅ {len(imported_data)}個のプロンプトをインポートしました"
                    )
                    st.rerun()
            except json.JSONDecodeError:
                st.error("❌ 無効なJSONファイルです")


def show_main_page():
    """メインページを表示"""
    st.title("📝 AI議事録作成ツール")
    st.markdown(
        """
    **🎤 音声/動画ファイル** → Whisperで文字起こし → Geminiで議事録生成  
    **📝 テキスト/Word/VTT** → 直接入力またはペースト → Geminiで議事録生成
    
    ---
    """
    )

    # API Key確認
    if not st.session_state.api_key:
        st.warning("⚠️ Gemini API Keyが設定されていません")
        st.info("👈 サイドバーの「設定」からAPI Keyを設定してください")

    # サイドバー設定
    st.sidebar.title("⚙️ Whisper設定")

    # ファイルアップロード
    st.markdown("### 📂 ファイルをアップロード")
    
    # 重要な注意書き
    st.info("""
    **💡 ファイルアップロード方法:**
    
    **📁 対応形式（そのままアップロード可能）**
    - 🎤 音声/動画: MP3, WAV, M4A, OGG, FLAC, MP4, AVI, MOV, MKV
    - 📝 Word: DOCX, DOC
    - 🎬 字幕: VTT
    - 📄 テキスト: TXT, MD
    
    **⚠️ エラーが出る場合の対処法:**
    - `.docx` → `.doc.txt` にリネーム
    - `.vtt` → `.vtt.txt` にリネーム
    - または「📝 テキスト直接入力」タブを使用
    """)
    
    # タブで入力方法を選択
    input_tab1, input_tab2 = st.tabs(["📁 ファイルアップロード", "📝 テキスト直接入力"])
    
    with input_tab1:
        uploaded_file = st.file_uploader(
            "ファイルを選択",
            type=[
                # 音声・動画
                "mp3", "wav", "m4a", "ogg", "flac",
                "mp4", "avi", "mov", "mkv",
                # テキスト
                "txt", "md", "text",
                # Word
                "docx", "doc",
                # VTT
                "vtt",
            ],
            help="対応形式: 音声/動画/テキスト/Word/VTT",
        )
        
        # ファイル名から元の形式を推測
        if uploaded_file:
            filename = uploaded_file.name.lower()
            if ".doc.txt" in filename or ".docx.txt" in filename:
                st.info("📝 Wordファイル（リネーム版）として処理します")
                st.session_state.file_type = "word_renamed"
            elif ".vtt.txt" in filename:
                st.info("🎬 VTTファイル（リネーム版）として処理します")
                st.session_state.file_type = "vtt_renamed"
            elif filename.endswith(('.docx', '.doc')):
                st.info("📝 Wordファイルとして処理します")
                st.session_state.file_type = "word"
            elif filename.endswith('.vtt'):
                st.info("🎬 VTTファイルとして処理します")
                st.session_state.file_type = "vtt"
    
    with input_tab2:
        st.markdown("""
        **💡 Word/VTTファイルの使い方:**
        1. Wordファイルを開いて全文コピー（Ctrl+A → Ctrl+C）
        2. VTTファイルをテキストエディタで開いてコピー
        3. 下のテキストエリアに貼り付け
        """)
        
        pasted_text = st.text_area(
            "テキストを貼り付け",
            height=300,
            placeholder="Word/VTT/テキストファイルの内容をここに貼り付けてください...",
            key="pasted_text_input"
        )
        
        if pasted_text:
            # VTT形式かどうかを判定
            is_vtt_format = pasted_text.strip().startswith("WEBVTT")
            
            if st.button("📝 このテキストを使用", type="primary"):
                if is_vtt_format:
                    st.info("🎬 VTT形式を検出しました。字幕テキストを抽出します...")
                    try:
                        # VTT形式のテキストを解析
                        text_lines = []
                        lines = pasted_text.split("\n")
                        
                        in_cue = False
                        for line in lines:
                            line = line.strip()
                            
                            # タイムスタンプ行をスキップ
                            if "-->" in line:
                                in_cue = True
                                continue
                            
                            # 空行はキューの終わり
                            if not line:
                                in_cue = False
                                continue
                            
                            # キュー識別子（数字のみの行）をスキップ
                            if line.isdigit():
                                continue
                            
                            # WEBVTTヘッダーやNOTE、STYLEなどをスキップ
                            if line.startswith(("WEBVTT", "NOTE", "STYLE", "REGION")):
                                continue
                            
                            # テキスト行を追加
                            if in_cue:
                                # VTTタグを削除（<c>、<v>など）
                                clean_line = re.sub(r'<[^>]+>', '', line)
                                if clean_line:
                                    text_lines.append(clean_line)
                        
                        extracted_text = " ".join(text_lines)
                        st.session_state.transcribed_text = extracted_text
                        st.success(f"✅ VTT字幕からテキストを抽出しました（{len(extracted_text)}文字）")
                    except Exception as e:
                        st.error(f"❌ VTT解析エラー: {str(e)}")
                else:
                    st.session_state.transcribed_text = pasted_text
                    st.success(f"✅ テキストを読み込みました（{len(pasted_text)}文字）")
                
                st.rerun()

    if uploaded_file is None and not st.session_state.transcribed_text:
        # 既存の文字起こしテキストがある場合は表示
        if st.session_state.transcribed_text:
            st.success("✅ 文字起こしテキストが保存されています")
            st.text_area(
                "保存済みテキスト",
                value=st.session_state.transcribed_text,
                height=200,
                disabled=True,
            )

            col1, col2 = st.columns([1, 4])
            with col1:
                if st.button("🗑️ クリア", type="secondary"):
                    st.session_state.transcribed_text = ""
                    st.session_state.minutes = ""
                    st.rerun()
        else:
            st.info("👆 ファイルをアップロードしてください")

            # 説明
            st.markdown("---")
            st.markdown("### 💡 使い方")
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(
                    """
                #### 🎤 音声/動画ファイル
                - MP3, WAV, M4A, OGG, FLAC
                - MP4, AVI, MOV, MKV
                - 自動で文字起こし → 議事録生成
                """
                )
            with col2:
                st.markdown(
                    """
                #### 📝 テキスト/Word/VTT
                - TXTファイルは直接アップロード
                - Word/VTTは「テキスト直接入力」タブからペースト
                - 議事録を直接生成
                """
                )

    if uploaded_file:
        # ファイルタイプを判定
        filename = uploaded_file.name

        # リネームされたWord/VTTファイルの処理
        if ".doc.txt" in filename.lower() or ".docx.txt" in filename.lower():
            # ==================================
            # リネームされたWordファイルの処理
            # ==================================
            st.session_state.file_type = "word_renamed"
            st.success("📝 Wordファイル（リネーム版）を処理します")
            
            try:
                # TXTとして読み込んでWordとして解析を試みる
                with st.spinner("📄 Wordファイルからテキストを抽出中..."):
                    # まずはWordファイルとして読み込みを試みる
                    try:
                        text_content = read_word_file(uploaded_file)
                        st.success(f"✅ Wordファイルからテキストを抽出しました（{len(text_content)}文字）")
                    except:
                        # Wordとして読めない場合はテキストとして処理
                        text_content = uploaded_file.read().decode("utf-8")
                        st.info("💡 Wordファイルとして読み込めなかったため、テキストとして処理しました")
                    
                    st.session_state.transcribed_text = text_content

                st.markdown("---")
                st.markdown("### 📄 抽出したテキスト")
                st.text_area("内容", value=text_content, height=300, key="word_renamed_display")
                
            except Exception as e:
                st.error(f"❌ {str(e)}")
                return
        
        elif ".vtt.txt" in filename.lower():
            # ==================================
            # リネームされたVTTファイルの処理
            # ==================================
            st.session_state.file_type = "vtt_renamed"
            st.success("🎬 VTTファイル（リネーム版）を処理します")
            
            try:
                with st.spinner("📄 VTTファイルからテキストを抽出中..."):
                    text_content = read_vtt_file(uploaded_file)
                    st.session_state.transcribed_text = text_content

                st.markdown("---")
                st.markdown("### 📄 抽出したテキスト")
                st.text_area("内容", value=text_content, height=300, key="vtt_renamed_display")
                st.success(f"✅ VTTファイルからテキストを抽出しました（{len(text_content)}文字）")
                
            except Exception as e:
                st.error(f"❌ {str(e)}")
                return

        elif is_audio_or_video_file(filename):
            # ==================================
            # 音声/動画ファイルの処理
            # ==================================
            st.session_state.file_type = "audio_video"

            st.success("🎤 音声/動画ファイルを検出しました")
            st.info("📝 Whisperで文字起こしを実行します")

            # Whisper設定
            model_option = st.sidebar.selectbox(
                "モデルサイズ",
                get_available_models(),
                index=1,
                help="小さいモデルほど高速ですが精度は低下します",
            )

            model_info = {
                "tiny": "最速・最小（39M）",
                "base": "高速・小型（74M）",
                "small": "バランス型（244M）",
                "medium": "高精度（769M）",
            }
            st.sidebar.caption(f"💡 {model_info.get(model_option, '')}")

            language_option = st.sidebar.selectbox(
                "言語",
                options=["", "en", "ja", "zh", "de", "fr", "es", "ko", "ru"],
                index=2,  # デフォルト日本語
                format_func=lambda x: {
                    "": "🌐 自動検出",
                    "en": "🇺🇸 英語",
                    "ja": "🇯🇵 日本語",
                    "zh": "🇨🇳 中国語",
                    "de": "🇩🇪 ドイツ語",
                    "fr": "🇫🇷 フランス語",
                    "es": "🇪🇸 スペイン語",
                    "ko": "🇰🇷 韓国語",
                    "ru": "🇷🇺 ロシア語",
                }.get(x, x),
            )

            device = "GPU (CUDA)" if torch.cuda.is_available() else "CPU"
            st.sidebar.info(f"🖥️ デバイス: **{device}**")

            # FFmpegチェック
            if not check_ffmpeg():
                st.error("⚠️ FFmpegが利用できません。")
                st.stop()

            st.markdown("---")

            # 文字起こし実行
            if st.button("🚀 文字起こし開始", type="primary", use_container_width=True):
                transcribed_text = transcribe_audio(
                    uploaded_file, model_option, language_option
                )

                if transcribed_text:
                    st.session_state.transcribed_text = transcribed_text

                    st.markdown("---")
                    st.markdown("### 📄 文字起こし結果")

                    st.text_area(
                        "テキスト",
                        value=st.session_state.transcribed_text,
                        height=300,
                        key="transcript_display",
                    )

                    st.download_button(
                        label="💾 文字起こしテキストをダウンロード",
                        data=st.session_state.transcribed_text,
                        file_name=f"{os.path.splitext(uploaded_file.name)[0]}_transcript.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )

                    st.success(
                        "✅ 文字起こし完了！下にスクロールして議事録を生成してください"
                    )

        elif is_word_file(filename):
            # ==================================
            # 通常のWordファイルの処理
            # ==================================
            st.session_state.file_type = "word"

            st.success("📝 Wordファイルを検出しました")
            st.info("📖 テキストを抽出してGeminiで議事録を生成します")

            # Wordファイル読み込み
            try:
                with st.spinner("📄 Wordファイルからテキストを抽出中..."):
                    text_content = read_word_file(uploaded_file)
                    st.session_state.transcribed_text = text_content

                st.markdown("---")
                st.markdown("### 📄 抽出したテキスト")

                st.text_area(
                    "内容", value=text_content, height=300, key="word_text_display"
                )

                st.success(f"✅ Wordファイルからテキストを抽出しました（{len(text_content)}文字）")

            except Exception as e:
                st.error(f"❌ {str(e)}")
                if not DOCX_AVAILABLE:
                    st.info("💡 ターミナルで以下を実行してください:\n```\npip install python-docx\n```")
                return

        elif is_text_file(filename):
            # ==================================
            # テキストファイルの処理（VTT含む）
            # ==================================
            st.session_state.file_type = "text"

            # VTTファイルかどうかをチェック
            if is_vtt_file(filename):
                st.success("🎬 VTT字幕ファイルを検出しました")
                st.info("📝 字幕テキストを抽出してGeminiで議事録を生成します")
                
                try:
                    with st.spinner("📄 VTTファイルからテキストを抽出中..."):
                        text_content = read_vtt_file(uploaded_file)
                        st.session_state.transcribed_text = text_content

                    st.markdown("---")
                    st.markdown("### 📄 抽出したテキスト")

                    st.text_area(
                        "内容", value=text_content, height=300, key="vtt_text_display"
                    )

                    st.success(f"✅ VTTファイルからテキストを抽出しました（{len(text_content)}文字）")

                except Exception as e:
                    st.error(f"❌ {str(e)}")
                    return
            else:
                st.success("📄 テキストファイルを検出しました")
                st.info("🤖 直接Geminiで議事録を生成します")

                # テキスト読み込み
                try:
                    text_content = uploaded_file.read().decode("utf-8")
                    st.session_state.transcribed_text = text_content

                    st.markdown("---")
                    st.markdown("### 📄 読み込んだテキスト")

                    st.text_area(
                        "内容", value=text_content, height=300, key="loaded_text_display"
                    )

                    st.success(f"✅ テキストを読み込みました（{len(text_content)}文字）")

                except Exception as e:
                    st.error(f"❌ テキストの読み込みに失敗しました: {str(e)}")
                    return

        elif is_word_file(filename):
            # ==================================
            # Wordファイルの処理
            # ==================================
            st.session_state.file_type = "word"

            st.success("📝 Wordファイルを検出しました")
            st.info("📖 テキストを抽出してGeminiで議事録を生成します")

            # Wordファイル読み込み
            try:
                with st.spinner("📄 Wordファイルからテキストを抽出中..."):
                    text_content = read_word_file(uploaded_file)
                    st.session_state.transcribed_text = text_content

                st.markdown("---")
                st.markdown("### 📄 抽出したテキスト")

                st.text_area(
                    "内容", value=text_content, height=300, key="word_text_display"
                )

                st.success(f"✅ Wordファイルからテキストを抽出しました（{len(text_content)}文字）")

            except Exception as e:
                st.error(f"❌ {str(e)}")
                if not DOCX_AVAILABLE:
                    st.info("💡 ターミナルで以下を実行してください:\n```\npip install python-docx\n```")
                return

        elif is_vtt_file(filename):
            # ==================================
            # 通常のVTTファイルの処理
            # ==================================
            st.session_state.file_type = "vtt"

            st.success("🎬 VTTファイルを検出しました")
            st.info("📝 字幕テキストを抽出してGeminiで議事録を生成します")

            # VTTファイル読み込み
            try:
                with st.spinner("📄 VTTファイルからテキストを抽出中..."):
                    text_content = read_vtt_file(uploaded_file)
                    st.session_state.transcribed_text = text_content

                st.markdown("---")
                st.markdown("### 📄 抽出したテキスト")

                st.text_area(
                    "内容", value=text_content, height=300, key="vtt_direct_text_display"
                )

                st.success(f"✅ VTTファイルからテキストを抽出しました（{len(text_content)}文字）")

            except Exception as e:
                st.error(f"❌ {str(e)}")
                if not WEBVTT_AVAILABLE:
                    st.info("💡 webvtt-pyライブラリのインストールを推奨します")
                return

        else:
            st.error("❌ サポートされていないファイル形式です")
            return

    # ==================================
    # 議事録生成セクション（共通）
    # ==================================
    if st.session_state.transcribed_text:
        st.markdown("---")
        st.markdown("## 🤖 議事録を生成")

        # API Key確認
        if not st.session_state.api_key:
            st.error("⚠️ Gemini API Keyが設定されていません")
            st.info("👈 サイドバーの「設定」からAPI Keyを設定してください")
            return

        # 現在のモデル表示
        st.info(f"🤖 使用モデル: **{st.session_state.gemini_model}**")

        # プロンプトテンプレート選択
        st.markdown("### 📋 議事録のスタイル選択")

        templates = get_all_prompt_templates()
        template_names = list(templates.keys())

        selected_template_name = st.selectbox(
            "テンプレート", template_names, help="用途に応じて選択してください"
        )

        # プレビュー表示
        with st.expander("📝 プロンプトプレビュー"):
            st.code(templates[selected_template_name], language="markdown")

        # 議事録生成ボタン
        st.markdown("---")
        if st.button("✨ 議事録を生成", type="primary", use_container_width=True):
            with st.spinner("🤖 Geminiが議事録を生成中..."):
                try:
                    start_time = time.time()

                    minutes = generate_minutes_with_gemini(
                        st.session_state.transcribed_text,
                        templates[selected_template_name],
                        st.session_state.api_key,
                        st.session_state.gemini_model,
                    )

                    generation_time = time.time() - start_time
                    st.session_state.minutes = minutes

                    st.success(f"✅ 議事録生成完了！（{generation_time:.2f}秒）")

                except Exception as e:
                    st.error(f"❌ エラー: {str(e)}")
                    st.info("💡 API Keyが正しいか、設定ページで確認してください")
                    return

        # 生成された議事録を表示
        if st.session_state.minutes:
            st.markdown("---")
            st.markdown("## 📄 生成された議事録")

            # タブで表示切り替え
            view_tab1, view_tab2 = st.tabs(["📖 プレビュー", "📝 編集"])

            with view_tab1:
                st.markdown(st.session_state.minutes)

            with view_tab2:
                edited_minutes = st.text_area(
                    "議事録を編集",
                    value=st.session_state.minutes,
                    height=500,
                    key="edit_minutes",
                )

                if st.button("💾 編集を保存", type="secondary"):
                    st.session_state.minutes = edited_minutes
                    st.success("✅ 保存しました")

            # ダウンロードボタン
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📥 Markdownでダウンロード",
                    data=st.session_state.minutes,
                    file_name=f"議事録_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

            with col2:
                st.download_button(
                    label="📥 テキストでダウンロード",
                    data=st.session_state.minutes,
                    file_name=f"議事録_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    use_container_width=True,
                )


def main():
    """メイン関数"""
    # サイドバーでページ選択
    page = st.sidebar.radio(
        "ナビゲーション", ["🏠 ホーム", "⚙️ 設定"], label_visibility="collapsed"
    )

    if page == "🏠 ホーム":
        show_main_page()
    else:
        show_settings_page()


if __name__ == "__main__":
    main()
